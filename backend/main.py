# backend/main.py
import os
import asyncio
import logging
import json
import pandas as pd
from datetime import datetime
import openai
from dotenv import load_dotenv
from fastapi import FastAPI, WebSocket, WebSocketDisconnect
from fastapi.middleware.cors import CORSMiddleware

from deepgram import (
    DeepgramClient,
    DeepgramClientOptions,
    LiveTranscriptionEvents,
    LiveOptions,
)

# --- Configuración Inicial ---
load_dotenv()  # Carga variables de entorno desde .env

API_KEY = os.getenv("DEEPGRAM_API_KEY")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not API_KEY:
    raise ValueError("DEEPGRAM_API_KEY no encontrada en las variables de entorno.")
if not OPENAI_API_KEY:
    raise ValueError("OPENAI_API_KEY no encontrada en las variables de entorno.")

# Configurar OpenAI
openai.api_key = OPENAI_API_KEY

# Setup logging
logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

# Deepgram Client Configuration
config: DeepgramClientOptions = DeepgramClientOptions(
    verbose=logging.DEBUG
)

# Initialize Deepgram Client
deepgram: DeepgramClient = DeepgramClient(API_KEY, config)

# Initialize FastAPI app - KEEP ONLY THIS INSTANCE
app = FastAPI()

# Configure CORS - MOVE THIS HERE
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allow all origins for now
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.websocket("/ws/transcribe")
async def websocket_endpoint(websocket: WebSocket):
    """Maneja conexiones WebSocket de clientes para transcripción en tiempo real."""
    await websocket.accept()
    logger.info(f"Cliente conectado: {websocket.client}")

    dg_connection = None
    full_transcript = []  # Almacena la transcripción completa
    is_final = False  # Indica si el fragmento es final o parcial
    
    # Default model configuration
    selected_model = "nova-2"  # Default model for 2-person conversations

    try:
        dg_connection = deepgram.listen.asynclive.v("1")
        logger.info("Conexión asíncrona de Deepgram creada.")

        async def on_message(self, result, **kwargs):
            """Callback para cuando se recibe una transcripción."""
            nonlocal full_transcript  # Add this line to access the outer variable
            
            try:
                transcript = result.channel.alternatives[0].transcript
                is_final = result.is_final  # Use is_final directly from the response
                
                if len(transcript) > 0:
                    # Solo enviamos al cliente los resultados finales o parciales con indicador
                    message = {
                        "transcript": transcript,
                        "is_final": is_final
                    }
                    
                    # Si es un resultado final, lo guardamos para análisis posterior
                    if is_final and transcript.strip():
                        logger.info(f"Adding final transcript segment: '{transcript}'")
                        full_transcript.append(transcript)
                        logger.info(f"Current transcript segments: {len(full_transcript)}")
                    
                    await websocket.send_text(json.dumps(message))
            except WebSocketDisconnect:
                logger.warning("Cliente desconectado antes de enviar mensaje.")
            except Exception as e:
                logger.error(f"Error al procesar/enviar mensaje: {e}")

        # --- Resto de handlers de eventos ---
        async def on_metadata(self, metadata, **kwargs):
            logger.debug(f"Deepgram metadata: {metadata}")

        async def on_speech_started(self, speech_started, **kwargs):
            logger.debug("Deepgram speech_started")

        async def on_utterance_end(self, utterance_end, **kwargs):
            logger.debug("Deepgram utterance_end")

        async def on_error(self, error, **kwargs):
            error_message = error.get("message", str(error))
            logger.error(f"Error de Deepgram: {error_message}")
            try:
                await websocket.send_text(json.dumps({"error": error_message}))
            except Exception as e:
                logger.error(f"Error al enviar mensaje de error: {e}")

        async def on_open(self, open_event, **kwargs):
            """Callback cuando la conexión con Deepgram se abre."""
            logger.info(f"Conexión Deepgram abierta: {open_event}")

        # Fix the on_close handler to properly handle the close_event parameter
        async def on_close(self, close_event=None, **kwargs):
            """Callback cuando la conexión con Deepgram se cierra."""
            logger.info(f"Conexión Deepgram cerrada: {close_event}")

        # --- Registrar Handlers de Eventos ---
        dg_connection.on(LiveTranscriptionEvents.Transcript, on_message)
        dg_connection.on(LiveTranscriptionEvents.Metadata, on_metadata)
        dg_connection.on(LiveTranscriptionEvents.SpeechStarted, on_speech_started)
        dg_connection.on(LiveTranscriptionEvents.UtteranceEnd, on_utterance_end)
        dg_connection.on(LiveTranscriptionEvents.Error, on_error)
        dg_connection.on(LiveTranscriptionEvents.Open, on_open)
        dg_connection.on(LiveTranscriptionEvents.Close, on_close)

        # Wait for initial configuration from client before starting
        logger.info("Esperando configuración inicial del cliente...")
        try:
            config_message_raw = await websocket.receive_text()
            config_message = json.loads(config_message_raw)
            
            # Check if client sent model configuration
            if "model" in config_message:
                selected_model = config_message["model"]
                logger.info(f"Modelo seleccionado por el cliente: {selected_model}")
            else:
                # Inform client about default model
                await websocket.send_text(json.dumps({
                    "status": "config",
                    "model": selected_model
                }))
        except Exception as e:
            logger.warning(f"Error al recibir configuración inicial, usando valores por defecto: {e}")
            # Inform client about default model
            await websocket.send_text(json.dumps({
                "status": "config",
                "model": selected_model
            }))

        # --- Opciones de Transcripción de Deepgram ---
        options = LiveOptions(
            model=selected_model,
            language="es",
            smart_format=True,
            interim_results=True,
            utterance_end_ms="1000",
            vad_events=True,
        )

        await dg_connection.start(options)
        logger.info(f"Conexión Deepgram iniciada con modelo {selected_model} y lista para recibir audio.")

        # --- Bucle Principal ---
        while True:
            try:
                # Recibir mensaje del cliente
                message_raw = await websocket.receive()
                
                # Determinar el tipo de mensaje
                if "type" in message_raw and message_raw["type"] == "websocket.receive":
                    if "text" in message_raw:
                        # Es un mensaje de texto (JSON)
                        try:
                            text_data = message_raw["text"]
                            logger.info(f"Received text message: {text_data}")
                            message = json.loads(text_data)
                            
                            # Si recibimos un comando para detener y analizar
                            if message.get("command") == "stop_and_analyze":
                                logger.info("Received stop_and_analyze command")
                                
                                # Get export format preferences
                                export_formats = message.get("export_formats", ["excel"])  # Default to Excel if not specified
                                logger.info(f"Requested export formats: {export_formats}")
                                
                                # Check if transcript was sent directly from frontend
                                client_transcript = message.get("transcript")
                                
                                if client_transcript:
                                    logger.info("Using transcript sent from client")
                                    complete_text = client_transcript
                                    analysis = await generate_analysis(complete_text)
                                    
                                    # Generate files based on requested formats
                                    file_data = {}
                                    
                                    try:
                                        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                                        
                                        if "excel" in export_formats:
                                            excel_data = await generate_excel_file(complete_text, analysis, timestamp)
                                            file_data["excel"] = {
                                                "filename": f"transcripcion_{timestamp}.xlsx",
                                                "data": excel_data,
                                                "content_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                                            }
                                        
                                        if "pdf" in export_formats:
                                            pdf_data = await generate_pdf_file(complete_text, analysis, timestamp)
                                            file_data["pdf"] = {
                                                "filename": f"transcripcion_{timestamp}.pdf",
                                                "data": pdf_data,
                                                "content_type": "application/pdf"
                                            }
                                        
                                        if "word" in export_formats:
                                            word_data = await generate_word_file(complete_text, analysis, timestamp)
                                            file_data["word"] = {
                                                "filename": f"transcripcion_{timestamp}.docx",
                                                "data": word_data,
                                                "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                            }
                                        
                                        # Send analysis and file data to client
                                        await websocket.send_text(json.dumps({
                                            "analysis_complete": True,
                                            "analysis": analysis,
                                            "file_data": {
                                                format_name: {
                                                    "filename": file_info["filename"],
                                                    "content_type": file_info["content_type"],
                                                    # We'll send the actual file data in a separate message
                                                }
                                                for format_name, file_info in file_data.items()
                                            }
                                        }))
                                        
                                        # Send each file separately to avoid large JSON messages
                                        for format_name, file_info in file_data.items():
                                            await websocket.send_bytes(file_info["data"])
                                            logger.info(f"Sent {format_name} file to client")
                                        
                                        logger.info(f"Análisis completado y archivos enviados al cliente")
                                    except Exception as e:
                                        logger.error(f"Error al generar archivos: {e}")
                                        # Enviar solo el análisis sin archivos
                                        await websocket.send_text(json.dumps({
                                            "analysis_complete": True,
                                            "analysis": analysis,
                                            "error_saving": str(e)
                                        }))
                                        logger.info("Análisis completado pero no se pudieron generar los archivos")
                                # If no client transcript, try to use the backend's stored transcript
                                elif full_transcript:
                                    logger.info(f"Analyzing transcript with {len(full_transcript)} segments")
                                    complete_text = " ".join(full_transcript)
                                    logger.info(f"Complete text for analysis: '{complete_text}'")
                                    analysis = await generate_analysis(complete_text)
                                    
                                    # Guardar en Excel
                                    filename = f"transcripcion_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
                                    filepath = os.path.join(os.getcwd(), filename)
                                    
                                    # Crear DataFrame y guardar
                                    df = pd.DataFrame({
                                        "Transcripción Completa": [complete_text],
                                        "Resumen General": [analysis.get("resumen", "")],
                                        "Percepciones por área": [analysis.get("percepciones_por_area", "")],
                                        "Relaciones entre áreas": [analysis.get("relaciones_entre_areas", "")],
                                        "Factores Experiencia": [analysis.get("factores_experiencia", "")],
                                        "Análisis de Sentimiento": [analysis.get("analisis_sentimiento", "")],
                                        "Recomendaciones": [analysis.get("recomendaciones", "")]
                                    })
                                    
                                    df.to_excel(filepath, index=False)
                                    
                                    # Enviar ruta del archivo al cliente
                                    await websocket.send_text(json.dumps({
                                        "analysis_complete": True,
                                        "file_path": filepath,
                                        "analysis": analysis
                                    }))
                                    
                                    logger.info(f"Análisis completado y guardado en {filepath}")
                                else:
                                    logger.warning("No transcript data to analyze")
                                    await websocket.send_text(json.dumps({
                                        "error": "No hay transcripción para analizar"
                                    }))
                        except json.JSONDecodeError as e:
                            logger.error(f"Error al decodificar mensaje JSON del cliente: {e}")
                            logger.error(f"Mensaje recibido: {message_raw}")
                    
                    elif "bytes" in message_raw:
                        # Es un mensaje de bytes (audio)
                        audio_data = message_raw["bytes"]
                        await dg_connection.send(audio_data)
                else:
                    logger.warning(f"Mensaje no reconocido: {message_raw}")
                
            except WebSocketDisconnect:
                logger.info(f"Cliente desconectado: {websocket.client}")
                break
            except Exception as e:
                logger.error(f"Error en el bucle principal: {e}")
                logger.exception("Detalle del error:")  # Esto imprimirá el stack trace completo
                break

    except Exception as e:
        logger.error(f"Error durante la configuración: {e}")
        try:
            await websocket.close(code=1011, reason=f"Error interno: {e}")
        except Exception:
            pass
    finally:
        if dg_connection:
            await dg_connection.finish()
            logger.info("Conexión Deepgram cerrada.")
        logger.info(f"Limpieza completa para cliente: {websocket.client}")

async def generate_analysis(text):
    """Genera análisis de la transcripción usando ChatGPT."""
    try:
        logger.info(f"Generando análisis para texto de {len(text)} caracteres")
        
        # Usar la API de OpenAI correctamente (versión async)
        from openai import AsyncOpenAI
        
        # Inicializar el cliente
        client = AsyncOpenAI(api_key=OPENAI_API_KEY)
        
        # Nuevo prompt para análisis organizacional
        messages = [
            {
                "role": "system",
                "content": "Eres un analista organizacional experto en experiencia del empleado, clima laboral y relaciones interdepartamentales. Tu tarea es analizar transcripciones de entrevistas a profundidad con empleados de una empresa, con el objetivo de identificar percepciones por área, relaciones entre áreas y oportunidades de mejora en la experiencia del empleado. No separes el texto por oradores ni intentes identificar quién habla. Analiza todo el contenido como un texto continuo."
            },
            {
                "role": "user",
                "content": f"""Analiza la siguiente transcripción y entrega un informe estructurado que contenga:
         
1. **Resumen general** de los temas tratados.
2. **Percepciones por área** (identifica cada área mencionada y describe lo que se dice sobre ella).
3. **Relaciones entre áreas** (cómo se percibe la colaboración, comunicación o fricciones entre áreas).
4. **Factores que afectan positiva o negativamente la experiencia del empleado**.
5. **Análisis de sentimiento general** y por cada área mencionada.
6. **Recomendaciones accionables** para mejorar la experiencia del empleado, basadas en lo expresado en la entrevista.
         
Transcripción: {text}
"""
            }
        ]
        
        # Realizar la solicitud con el nuevo prompt
        response = await client.chat.completions.create(
            model="gpt-4o",
            messages=messages
        )
        
        analysis_text = response.choices[0].message.content
        logger.info("Análisis generado correctamente")
        
        # Intentar estructurar la respuesta
        analysis = {
            "resumen": "",
            "percepciones_por_area": "",
            "relaciones_entre_areas": "",
            "factores_experiencia": "",
            "analisis_sentimiento": "",
            "recomendaciones": "",
            "texto_completo": analysis_text
        }
        
        # Extraer secciones (implementación simple)
        sections = analysis_text.split("\n\n")
        if len(sections) >= 6:
            analysis["resumen"] = sections[0].replace("Resumen general:", "").strip()
            analysis["percepciones_por_area"] = sections[1].replace("Percepciones por área:", "").strip()
            analysis["relaciones_entre_areas"] = sections[2].replace("Relaciones entre áreas:", "").strip()
            analysis["factores_experiencia"] = sections[3].replace("Factores que afectan positiva o negativamente la experiencia del empleado:", "").strip()
            analysis["analisis_sentimiento"] = sections[4].replace("Análisis de sentimiento general:", "").strip()
            analysis["recomendaciones"] = sections[5].replace("Recomendaciones accionables:", "").strip()
        
        return analysis
    except Exception as e:
        logger.error(f"Error al generar análisis con ChatGPT: {e}")
        logger.exception("Detalle del error:")
        return {"error": str(e), "texto_completo": text}

# Add these functions after the generate_analysis function

async def generate_excel_file(text, analysis, timestamp):
    """Genera un archivo Excel con el análisis."""
    try:
        # Crear DataFrame
        df = pd.DataFrame({
            "Transcripción Completa": [text],
            "Resumen General": [analysis.get("resumen", "")],
            "Percepciones por Área": [analysis.get("percepciones_por_area", "")],
            "Relaciones entre Áreas": [analysis.get("relaciones_entre_areas", "")],
            "Factores Experiencia": [analysis.get("factores_experiencia", "")],
            "Análisis de Sentimiento": [analysis.get("analisis_sentimiento", "")],
            "Recomendaciones": [analysis.get("recomendaciones", "")]
        })
        
        # Guardar en un buffer en memoria
        output = BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False)
        
        output.seek(0)
        return output.getvalue()
    except Exception as e:
        logger.error(f"Error al generar archivo Excel: {e}")
        raise

async def generate_pdf_file(text, analysis, timestamp):
    """Genera un archivo PDF con el análisis."""
    try:
        # Crear un buffer en memoria
        buffer = BytesIO()
        
        # Crear el documento PDF
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Título
        title_style = styles["Title"]
        story.append(Paragraph("Análisis de Transcripción", title_style))
        story.append(Spacer(1, 12))
        
        # Estilos para secciones y contenido
        heading_style = styles["Heading2"]
        normal_style = styles["Normal"]
        
        # Añadir secciones
        story.append(Paragraph("Resumen General", heading_style))
        story.append(Paragraph(analysis.get("resumen", ""), normal_style))
        story.append(Spacer(1, 12))
        
        story.append(Paragraph("Percepciones por Área", heading_style))
        story.append(Paragraph(analysis.get("percepciones_por_area", ""), normal_style))
        story.append(Spacer(1, 12))
        
        story.append(Paragraph("Relaciones entre Áreas", heading_style))
        story.append(Paragraph(analysis.get("relaciones_entre_areas", ""), normal_style))
        story.append(Spacer(1, 12))
        
        story.append(Paragraph("Factores que Afectan la Experiencia del Empleado", heading_style))
        story.append(Paragraph(analysis.get("factores_experiencia", ""), normal_style))
        story.append(Spacer(1, 12))
        
        story.append(Paragraph("Análisis de Sentimiento", heading_style))
        story.append(Paragraph(analysis.get("analisis_sentimiento", ""), normal_style))
        story.append(Spacer(1, 12))
        
        story.append(Paragraph("Recomendaciones", heading_style))
        story.append(Paragraph(analysis.get("recomendaciones", ""), normal_style))
        story.append(Spacer(1, 12))
        
        story.append(Paragraph("Transcripción Completa", heading_style))
        story.append(Paragraph(text, normal_style))
        
        # Construir el documento
        doc.build(story)
        
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        logger.error(f"Error al generar archivo PDF: {e}")
        raise

async def generate_word_file(text, analysis, timestamp):
    """Genera un archivo Word con el análisis."""
    try:
        # Crear un documento Word
        doc = Document()
        
        # Título
        doc.add_heading("Análisis de Transcripción", 0)
        
        # Añadir secciones
        doc.add_heading("Resumen General", level=1)
        doc.add_paragraph(analysis.get("resumen", ""))
        
        doc.add_heading("Percepciones por Área", level=1)
        doc.add_paragraph(analysis.get("percepciones_por_area", ""))
        
        doc.add_heading("Relaciones entre Áreas", level=1)
        doc.add_paragraph(analysis.get("relaciones_entre_areas", ""))
        
        doc.add_heading("Factores que Afectan la Experiencia del Empleado", level=1)
        doc.add_paragraph(analysis.get("factores_experiencia", ""))
        
        doc.add_heading("Análisis de Sentimiento", level=1)
        doc.add_paragraph(analysis.get("analisis_sentimiento", ""))
        
        doc.add_heading("Recomendaciones", level=1)
        doc.add_paragraph(analysis.get("recomendaciones", ""))
        
        doc.add_heading("Transcripción Completa", level=1)
        doc.add_paragraph(text)
        
        # Guardar en un buffer en memoria
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        logger.error(f"Error al generar archivo Word: {e}")
        raise

@app.get("/")
async def root():
    """Root endpoint that returns basic API information."""
    return {
        "message": "API de Transcripción en Tiempo Real con FastAPI y Deepgram. Conéctate vía WebSocket a /ws/transcribe"
    }

# ... Add these imports at the top of the file
from fastapi.responses import FileResponse
from io import BytesIO
import docx
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
import tempfile

# --- Para Ejecutar Localmente (opcional) ---
# Se recomienda usar `uvicorn main:app --host 0.0.0.0 --port 8000 --reload`
# if __name__ == "__main__":
#     import uvicorn
#     logger.info("Iniciando servidor Uvicorn...")
#     uvicorn.run(app, host="0.0.0.0", port=8000)